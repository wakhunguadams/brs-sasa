"""
Process legislation documents and add them to the knowledge base
Specifically handles the Trust Administration Bill 2025
"""
import asyncio
from pathlib import Path
from docx import Document
from core.knowledge_base import knowledge_base
from utils.document_loader import TextChunker
from core.logger import setup_logger

logger = setup_logger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a .docx file"""
    try:
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""

async def process_legislation_documents():
    """Process all legislation documents and add to knowledge base"""
    try:
        logger.info("Starting legislation document processing...")
        
        # Initialize knowledge base
        await knowledge_base.initialize()
        
        # Process Trust Administration Bill
        legislation_dir = Path("legislation")
        trust_bill_path = legislation_dir / "Trust-Administration-Bill-2025.txt"
        
        if not trust_bill_path.exists():
            logger.error(f"Trust Administration Bill not found at {trust_bill_path}")
            return
        
        logger.info(f"Processing {trust_bill_path.name}...")
        
        # Read text file
        with open(trust_bill_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        if not text:
            logger.error("Failed to extract text from Trust Administration Bill")
            return
        
        logger.info(f"Extracted {len(text)} characters from {trust_bill_path.name}")
        
        # Chunk the document
        chunker = TextChunker(chunk_size=1000, chunk_overlap=200)
        chunks = chunker.chunk_text(text, trust_bill_path.name)
        
        logger.info(f"Created {len(chunks)} chunks from {trust_bill_path.name}")
        
        # Add metadata to mark as legislation
        for chunk in chunks:
            chunk['metadata']['type'] = 'legislation'
            chunk['metadata']['document_type'] = 'bill'
            chunk['metadata']['legislation_name'] = 'Trust Administration Bill 2025'
        
        # Add to knowledge base
        texts = [chunk['content'] for chunk in chunks]
        metadatas = [chunk['metadata'] for chunk in chunks]
        
        await knowledge_base.add_documents(texts, metadatas)
        
        logger.info(f"Successfully added {len(chunks)} chunks to knowledge base")
        logger.info("Legislation processing complete!")
        
        return len(chunks)
        
    except Exception as e:
        logger.error(f"Error processing legislation documents: {str(e)}")
        raise

if __name__ == "__main__":
    asyncio.run(process_legislation_documents())
